
from django.shortcuts import render
from django.http import HttpResponse
from .forms import CVUploadForm
from .models import CV
import os
import re
import PyPDF2
import docx
import textract
import xml.etree.ElementTree as ET
import pandas as pd
from django.core.files.storage import FileSystemStorage
from django.conf import settings
import logging

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def extract_pdf_text(pdf_path):
    text = ""
    try:

        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
        logger.debug(f"Extracted text from PDF {pdf_path}: {text[:500]}")
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
    return text

def extract_docx_text(docx_path):
    #doc = docx.Document(docx_path)
    text = ""
    try:
        doc = docx.Document(docx_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text  + "\n"
        logger.debug(f"Extracted text from DOCX {docx_path}: {text[:500]}")
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
    return text

def extract_doc_text(doc_path):
    text = ""
    try:
        text = textract.process(doc_path).decode('utf-8')
        logger.debug(f"Extractedtext from DOC{doc_path}:{text[:500]}")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOC: {e}")
        return ""


def extract_xml_text(xml_path):
    #tree = ET.parse(xml_path)

    #root = tree.getroot()
    text = ""
    try:
        tree = ET.parse(xml_path)
        root = tree.getroot()
        for elem in root.iter():
            if elem.text:
                text += elem.text + " "
        logger.debug(f"Extracted text from XML {xml_path}: {text[:500]}")
    except Exception as e:
        logger.error(f"Error extracting text from XML: {e}")
    return text
def is_valid_phone_number(number):
    digits_only = re.sub(r'\D','',number)
    return 10 <= len(digits_only) <= 14

def is_date_pattern(number):
    date_patterns = [ 
        r'\b\d{1,2}/\d{1,2}/\d{2,4}\b', # Matches dd/mm/yyyy, d/m/yy, etc.
        r'\b\d{1,2}-\d{1,2}-\d{2,4}\b', # Matches dd-mm-yyyy, d-m-yy, etc.
        r'\b\d{1,2} \w+ \d{2,4}\b', # Matches dd Month yyyy 
        r'\b\w+ \d{1,2}, \d{2,4}\b' # Matches Month dd, yyyy 
    ]
    for pattern in date_patterns: 
        if re.search(pattern, number): 
            return True 
    return False
    

def extract_contact_info(text):
    emails = re.findall(r'[\w\.-]+@[\w\.-]+', text)
    potential_numbers = re.findall(r'\+?[1-9][0-9 .\-\(\)]{7,}[0-9]', text)
    phone_numbers = []
    for number in potential_numbers :
        if is_valid_phone_number(number) and not is_date_pattern(number):
            digits_only = re.sub(r'\D','',number)
            if not(1900<= int(digits_only[:4]) <= 2100):
                phone_numbers.append(number)
    
    logger.debug(f"Extracted emails: {emails}")
    logger.debug(f"Extracted phone numbers: {phone_numbers}")
    return emails, phone_numbers

def process_file(file):
    #file_path = os.path.join(settings.MEDIA_ROOT,file.name)
    #with open(file_path, 'wb+') as destination:
    #    for chunk in file.chunks():
     #       destination.write(chunk)
    #text = ""
    fs  = FileSystemStorage()

    filename = fs.save(file.name, file)
    file_path = fs.path(filename)

    logger.debug(f"Processing file: {file_path} ")

    text = ""

    if file_path.endswith(".pdf"):
        text = extract_pdf_text(file_path)
    elif file_path.endswith(".docx"):
        text = extract_docx_text(file_path)
    elif file_path.endswith(".doc"):
        text = extract_doc_text(file_path)
    elif file_path.endswith(".xml"):
        text = extract_xml_text(file_path)
    else:
        logger.warning(f"Unsupported file type: {file_path}")
        return None,None,None,None
    if not text:
        logger.warning(f"No text extracted from file: {file_path}")
    
    emails, phone_numbers = extract_contact_info(text)
    logger.debug(f"Extracted emails: {emails}, phone numbers: {phone_numbers}")

    max_length = 255
    contact_numbers = ', '.join(phone_numbers)
    if len(contact_numbers) > max_length:
        contact_numbers = contact_numbers[:max_length]
    return filename, emails, phone_numbers,text

def upload_folder(request):
    if request.method == 'POST':
        logger.debug("POST request received")
        form = CVUploadForm(request.POST, request.FILES)
       
        if form.is_valid():
            files = request.FILES.getlist('files')
            data = []
            logger.debug("Files uploaded: %s",files)

            for file in files:

                #fs = FileSystemStorage()
                #filename= fs.save(file.name,file)
                filename, emails, phone_numbers, text = process_file(file)
                if filename:
                    CV.objects.create(
                        filename=filename,
                        email = ', '.join(emails),
                        contact_number = phone_numbers,
                        text=text
                    )
                    data.append({
                        "Filename": filename,
                        "Email":', '.join(emails),
                        "Contact Number":phone_numbers,
                        "Text": text
                    })
            logger.debug(f"Data collected for Excel: {data}")

            if data:
                df = pd.DataFrame(data)
                output_excel_path = os.path.join(settings.MEDIA_ROOT,'output_data.xlsx')
                try:                   
                    df.to_excel(output_excel_path, index=False)
                    logger.debug(f"Excel file created at: {output_excel_path}")

            #output_url = f"{settings.MEDIA_URL}output_data.xlsx"
                    return HttpResponse(f"Excel file generated. <a href='{settings.MEDIA_URL}output_data.xlsx' > Download Excel File</a>")
                except Exception as e:
                    logger.error(f"Error creating Excel files: {e}")
                    return HttpResponse("Error creating excel file.")
            else:
                logger.warning("No data extracted from uploaded files.")
                return HttpResponse("No data extracted from uploaded files.")
        else:
            logger.warning(f"Form is invalid : {form.errors}")
            return HttpResponse(f"Form is invalid : {form.errors}")
                
    else:
        form = CVUploadForm()
    return render(request, 'upload.html',{'form':form})
   






















"""
    data = []
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        if os.path.isfile(file_path):
            cv_text = ""
            if filename.endswith(".pdf"):
                cv_text = extract_pdf_text(file_path)
            elif filename.endswith(".docx"):
                cv_text = extract_docx_text(file_path)
            elif filename.endswith(".xml"):
                cv_text = extract_xml_text(file_path)
            
            if cv_text:
                cv_emails, cv_phone_numbers = extract_contact_info(cv_text)
                cv = CV(
                    filename=filename,
                    email=', '.join(cv_emails),
                    contact_number=', '.join(cv_phone_numbers),
                    text=cv_text
                )
                data.append(cv)
    CV.objects.bulk_create(data)
    """
""" 

def upload_cv(request):
    if request.method == 'POST':
        form = CVUploadForm(request.POST, request.FILES)
        if form.is_valid():
            folder = request.FILES['folder']
            folder_path = os.path.join('media', folder.name)
            os.makedirs(folder_path, exist_ok=True)
            
            for uploaded_file in request.FILES.getlist('folder'):
                file_path = os.path.join(folder_path, uploaded_file.name)
                with open(file_path, 'wb') as destination:
                    for chunk in uploaded_file.chunks():
                        destination.write(chunk)
            
            process_folder(folder_path)
            return HttpResponse("CVs uploaded successfully!")
    else:
        form = CVUploadForm()
    return render(request, 'upload_cv.html', {'form': form})



import shutil
import zipfile


def extract_zip(zip_file, extract_to):
    with zipfile.ZipFile(zip_file, 'r') as zip_ref:
        zip_ref.extractall(extract_to)
        """
"""
def upload_folder(request):
    if request.method == 'POST':
        form = CVUploadForm(request.POST, request.FILES)
       
        if form.is_valid():
            files = request.FILES.getlist('folder')
            data = []
            for file in files:
                filename, emails, phone_numbers,text=process_file(file)
                if filename:
                    cv = CV.objects.create(
                        filename=filename,
                        email = ', '.join(emails),
                        contact_number = ', '.join(phone_numbers),
                        text=text
                    )
                    data.append({
                        "Filename": filename,
                        "Email":', '.join(emails),
                        "Contact Number":', '.join(phone_numbers),
                        "Text": text
                    })
            df = pd.DataFrame(data)
            output_excel_path = os.path.join(settings.MEDIA_ROOT,'output_data.xlsx')
            df.to_excel(output_excel_path, index=False)

            return HttpResponse(f"Excel file generated. <a href='/media/output_data.xlsx' > Download Excel File</a>")
        else:
            form = CVUploadForm()
        return render(request, 'upload_cv.html',{'form':form})
"""

"""
            folder_name = folder.name
            folder_path = os.path.join('uploads', folder_name)
            os.makedirs(folder_path, exist_ok=True)
            with open(folder_path, 'wb+') as destination:
                for chunk in folder.chunks():
                    destination.write(chunk)
            
            # Check if the uploaded file is a zip file
            if folder_name.endswith('.zip'):
                extract_to = os.path.join('uploads', folder_name.split('.')[0])
                extract_zip(folder_path, extract_to)
                folder_path = extract_to
            data = []
            
            # Process the files in the folder
            for filename in os.listdir(folder_path):
                file_path = os.path.join(folder_path, filename)
                # Process each file using your existing logic
                # Example:
                if filename.endswith(".pdf"):
                    text = extract_pdf_text(file_path)
                elif filename.endswith(".docx"):
                    text = extract_docx_text(file_path)
                elif filename.endswith(".xml"):
                    text = extract_xml_text(file_path)
                else:
                    continue

                emails, phone_numbers = extract_contact_info(text)

                cv_data = {
                    "Filename": filename,
                    "Email": emails,
                    "Phone Number": phone_numbers,
                    "Text": text
                }
                data.append(cv_data)

            # Create DataFrame and export to Excel
            df = pd.DataFrame(data)
            output_excel_path = os.path.join('downloads', 'output_data.xlsx')
            df.to_excel(output_excel_path, index=False)

            return HttpResponse("Excel file generated. <a href='/media/output_data.xlsx'>Download Excel file</a>")

    else:
        form

"""